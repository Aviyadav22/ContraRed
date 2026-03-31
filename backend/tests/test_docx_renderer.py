import pytest, io
from docx import Document as DocxDocument
from app.services.drafting.models import (
    RawDraft, DraftSection, DraftMetadata, FinalDraft, QualityReport,
)

def _make_final_draft():
    return FinalDraft(
        draft=RawDraft(
            contract_type="nda_mutual", title="MUTUAL NON-DISCLOSURE AGREEMENT",
            sections=[
                DraftSection(number="1", heading="Preamble", content="This Agreement is between Acme Inc. and Beta LLC.",
                             clause_type="preamble", tier_used="acceptable"),
                DraftSection(number="2", heading="Confidential Information",
                             content='"Confidential Information" means all non-public info.',
                             clause_type="confidential_info_definition", tier_used="acceptable"),
            ],
            defined_terms={"Confidential Information": "all non-public info"},
            metadata=DraftMetadata(playbook_id="test", model="test", generation_seconds=5, tokens_used=1000),
        ),
        quality_report=QualityReport(overall_score=90, risk_alignment=92, compliance_score=88,
                                      qa_score=90, annotations_applied=1, conflicts_flagged=0, open_annotations=[]),
    )

def test_render_docx_returns_bytes():
    from app.services.drafting.renderer.docx_renderer import render_docx
    result = render_docx(_make_final_draft())
    assert isinstance(result, bytes) and len(result) > 0

def test_render_docx_is_valid_document():
    from app.services.drafting.renderer.docx_renderer import render_docx
    data = render_docx(_make_final_draft())
    doc = DocxDocument(io.BytesIO(data))
    assert len(doc.paragraphs) > 0

def test_render_docx_has_title():
    from app.services.drafting.renderer.docx_renderer import render_docx
    data = render_docx(_make_final_draft())
    doc = DocxDocument(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("MUTUAL NON-DISCLOSURE AGREEMENT" in t for t in texts)

def test_render_docx_has_sections():
    from app.services.drafting.renderer.docx_renderer import render_docx
    data = render_docx(_make_final_draft())
    doc = DocxDocument(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("Preamble" in t for t in texts)
    assert any("Confidential Information" in t for t in texts)
