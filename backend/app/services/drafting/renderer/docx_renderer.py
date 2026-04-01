import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.drafting.models import FinalDraft


def render_docx(final_draft: FinalDraft) -> bytes:
    doc = Document()
    # Page setup: 1 inch margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    # Default font: Times New Roman 11pt, 1.15 spacing
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    # Title centered, bold, 16pt
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(final_draft.draft.title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    doc.add_paragraph()  # spacer
    # Sections
    for ds in final_draft.draft.sections:
        heading_text = f"{ds.number}. {ds.heading}" if ds.number else ds.heading
        heading = doc.add_heading(heading_text, level=2)
        for r in heading.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        for para_text in ds.content.strip().split("\n\n"):
            for line in para_text.strip().split("\n"):
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(4)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
