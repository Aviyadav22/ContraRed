"""
Batch Report DOCX Generator — produces a professional Word document
from batch analysis results for lawyer/client consumption.
"""

import io
import logging
from datetime import datetime, timezone
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# Brand colors
RED = RGBColor(0xEF, 0x44, 0x44)
YELLOW = RGBColor(0xF5, 0x9E, 0x0B)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
DARK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, color_hex: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def _risk_color(level: str) -> RGBColor:
    level = (level or "").lower()
    if level == "red":
        return RED
    if level == "yellow":
        return YELLOW
    return GREEN


def _risk_bg_hex(level: str) -> str:
    level = (level or "").lower()
    if level == "red":
        return "FEE2E2"
    if level == "yellow":
        return "FEF3C7"
    return "DCFCE7"


def generate_batch_report_docx(report_data: Dict[str, Any]) -> bytes:
    """Generate a professional Word document from batch analysis data.

    Args:
        report_data: The full batch report dict (from /batch/{id}/full-report)

    Returns:
        Bytes of the .docx file
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK

    # ---- COVER / TITLE ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    run = p.add_run('CONTRARED')
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT
    run.font.bold = True
    run.font.name = 'Calibri'

    title = doc.add_heading('Contract Portfolio Risk Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Generated: {datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC")}')
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{report_data.get("total_files", 0)} contracts analyzed')
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.add_paragraph()  # spacing

    # ---- EXECUTIVE SUMMARY ----
    doc.add_heading('Executive Summary', level=1)

    agg = report_data.get("aggregate_risk_summary", {})
    score = report_data.get("compliance_score", 0)

    # Summary table
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    headers = ['Compliance Score', 'Critical (Red)', 'Warnings (Yellow)', 'Deal Breakers', 'Total Findings']
    values = [
        f'{score}%',
        str(agg.get("red", 0)),
        str(agg.get("yellow", 0)),
        str(agg.get("deal_breakers", 0)),
        str(agg.get("total", 0)),
    ]

    for i, (header, value) in enumerate(zip(headers, values)):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = MUTED

        cell = table.rows[1].cells[i]
        cell.text = value
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                if i == 0:
                    run.font.color.rgb = GREEN if score >= 70 else YELLOW if score >= 40 else RED
                elif i in (1, 3):
                    run.font.color.rgb = RED
                elif i == 2:
                    run.font.color.rgb = YELLOW

    doc.add_paragraph()

    # ---- TOP RISKS ----
    top_risks = report_data.get("top_risks", [])
    if top_risks:
        doc.add_heading('Top Risks Across Portfolio', level=2)
        for i, risk in enumerate(top_risks[:10], 1):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(f'{risk["rule"]}')
            run.font.bold = True
            p.add_run(f'  ({risk["count"]} occurrences)')

        doc.add_paragraph()

    # ---- RISK HEATMAP ----
    heatmap = report_data.get("risk_heatmap", [])
    if heatmap:
        doc.add_heading('Risk Distribution by Clause Type', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Clause Type'
        table.rows[0].cells[1].text = 'Occurrences'
        for run in table.rows[0].cells[0].paragraphs[0].runs:
            run.font.bold = True
        for run in table.rows[0].cells[1].paragraphs[0].runs:
            run.font.bold = True

        for item in heatmap[:15]:
            row = table.add_row()
            row.cells[0].text = item["clause_type"]
            row.cells[1].text = str(item["count"])

        doc.add_paragraph()

    # ---- PER-CONTRACT ANALYSIS ----
    doc.add_heading('Contract-by-Contract Analysis', level=1)

    per_file = report_data.get("per_file", [])
    for file_idx, file_data in enumerate(per_file, 1):
        filename = file_data.get("filename", "Unknown")
        rs = file_data.get("risk_summary", {})
        findings = file_data.get("findings", [])
        ai_fallback = file_data.get("ai_fallback", False)

        # Contract header
        doc.add_heading(f'{file_idx}. {filename}', level=2)

        # Risk summary line
        p = doc.add_paragraph()
        if file_data.get("status") == "error":
            run = p.add_run(f'Error: {file_data.get("error", "Unknown error")}')
            run.font.color.rgb = RED
            continue

        run = p.add_run(f'{rs.get("red", 0)} Critical  |  ')
        run.font.color.rgb = RED
        run.font.bold = True
        run = p.add_run(f'{rs.get("yellow", 0)} Warnings  |  ')
        run.font.color.rgb = YELLOW
        run.font.bold = True
        green_count = rs.get("total", 0) - rs.get("red", 0) - rs.get("yellow", 0)
        run = p.add_run(f'{green_count} Compliant')
        run.font.color.rgb = GREEN
        run.font.bold = True

        if ai_fallback:
            p = doc.add_paragraph()
            run = p.add_run('Note: AI analysis was unavailable for this file. Results are based on rule-engine keyword matching only.')
            run.font.color.rgb = YELLOW
            run.font.italic = True
            run.font.size = Pt(9)

        # Findings table
        if findings:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.autofit = True

            # Headers
            for i, header in enumerate(['Risk', 'Clause', 'Issue', 'Recommended Fix']):
                cell = table.rows[0].cells[i]
                cell.text = header
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

            # Set column widths
            for row in table.rows:
                row.cells[0].width = Cm(2)
                row.cells[1].width = Cm(5.5)
                row.cells[2].width = Cm(4.5)
                row.cells[3].width = Cm(4.5)

            for finding in findings:
                row = table.add_row()
                risk_level = (finding.get("risk_level") or "yellow").upper()

                # Risk level cell
                cell = row.cells[0]
                p = cell.paragraphs[0]
                run = p.add_run(risk_level)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = _risk_color(risk_level)
                if finding.get("is_deal_breaker"):
                    p.add_run('\n')
                    run2 = p.add_run('DEAL BREAKER')
                    run2.font.size = Pt(7)
                    run2.font.bold = True
                    run2.font.color.rgb = RED
                _set_cell_bg(cell, _risk_bg_hex(risk_level))

                # Clause text
                cell = row.cells[1]
                clause = finding.get("clause_text", "")
                if len(clause) > 300:
                    clause = clause[:300] + "..."
                p = cell.paragraphs[0]
                run = p.add_run(clause)
                run.font.size = Pt(9)
                run.font.color.rgb = DARK

                # Rule name
                if finding.get("rule_name"):
                    p.add_run('\n')
                    run = p.add_run(finding["rule_name"])
                    run.font.size = Pt(8)
                    run.font.color.rgb = MUTED
                    run.font.italic = True

                # Explanation
                cell = row.cells[2]
                explanation = finding.get("ai_explanation") or "No explanation available"
                p = cell.paragraphs[0]
                run = p.add_run(explanation)
                run.font.size = Pt(9)

                # Fix
                cell = row.cells[3]
                fix = finding.get("suggested_fix") or "—"
                p = cell.paragraphs[0]
                run = p.add_run(fix)
                run.font.size = Pt(9)
                if fix != "—":
                    run.font.color.rgb = GREEN
        else:
            p = doc.add_paragraph()
            run = p.add_run('No findings for this contract.')
            run.font.color.rgb = GREEN
            run.font.italic = True

        doc.add_paragraph()  # spacing between contracts

    # ---- FOOTER ----
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This report was generated by ContraRed AI-Powered Contract Analysis.')
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.font.italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This is not legal advice. All findings should be reviewed by qualified legal counsel.')
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.font.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
